#!/usr/bin/env python3
"""
Generate Printable ICT Worksheets - Free Sheets System (Foi Volante)
Generates Word (.docx) documents for cls5-8 TIC lessons.

Usage:
    python tools/printables/generate_sheets.py              # All grades
    python tools/printables/generate_sheets.py --grade 8    # Single grade
    python tools/printables/generate_sheets.py --grade 8 --module 3  # Single module

Output: tools/printables/output/cls{N}_M{M}_{name}.docx
"""
import os
import re
import sys
import json
import html
import argparse
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# python-docx imports
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = Path(__file__).resolve().parent.parent.parent
CONTENT_DIR = BASE / 'content' / 'tic'
WORKSHEETS_DIR = BASE / 'data' / 'worksheets'
CONTENT_MAP = BASE / 'content_map.json'
OUTPUT_DIR = Path(__file__).resolve().parent / 'output'

# ─── Theory Extractor ─────────────────────────────────────────────────────────

class TheoryExtractor:
    """Extract theory text from lesson HTML files."""

    @staticmethod
    def extract_atoms(html_content):
        """Extract atom titles and content from lesson HTML."""
        atoms = []
        # Find all atom divs
        atom_pattern = re.compile(
            r'<div\s+class="atom"[^>]*>.*?'
            r'<h3\s+class="atom-title">(.*?)</h3>.*?'
            r'<div\s+class="atom-content">(.*?)</div>',
            re.DOTALL
        )
        for match in atom_pattern.finditer(html_content):
            title = match.group(1).strip()
            content_html = match.group(2).strip()
            # Clean HTML to plain text
            text = TheoryExtractor._html_to_text(content_html)
            atoms.append({'title': title, 'text': text})
        return atoms

    @staticmethod
    def extract_quizzes(html_content):
        """Extract quiz questions from data-quiz attributes."""
        quizzes = []
        # Pattern: data-quiz='[{"question": "...", "options": [...], "correct": "b"}]'
        quiz_pattern = re.compile(r"data-quiz='(\[.*?\])'", re.DOTALL)
        for match in quiz_pattern.finditer(html_content):
            try:
                quiz_data = json.loads(match.group(1))
                for q in quiz_data:
                    if q.get('question') and q.get('options'):
                        quizzes.append({
                            'question': html.unescape(q['question']),
                            'options': [html.unescape(opt) for opt in q['options']],
                            'correct': q.get('correct', 'a')
                        })
            except (json.JSONDecodeError, KeyError):
                continue
        return quizzes

    @staticmethod
    def extract_goal(html_content):
        """Extract lesson goal text."""
        match = re.search(r'class="goal-text"[^>]*>(.*?)</[^>]+>', html_content, re.DOTALL)
        if match:
            return TheoryExtractor._html_to_text(match.group(1)).strip('"').strip()
        return ""

    @staticmethod
    def extract_title(html_content):
        """Extract lesson title from h1."""
        match = re.search(r'<h1\s+class="lesson-title">(.*?)</h1>', html_content, re.DOTALL)
        if match:
            return TheoryExtractor._html_to_text(match.group(1)).strip()
        return ""

    @staticmethod
    def _html_to_text(html_str):
        """Convert HTML to clean plain text."""
        text = html_str
        # Extract <pre> code blocks first, preserve them with markers
        code_blocks = []
        def _save_code(m):
            code = m.group(1)
            # Clean inner HTML from code block
            code = re.sub(r'<[^>]+>', '', code)
            code = html.unescape(code).strip()
            idx = len(code_blocks)
            code_blocks.append(code)
            return f'\n[CODE_BLOCK_{idx}]\n'
        text = re.sub(r'<pre[^>]*>(.*?)</pre>', _save_code, text, flags=re.DOTALL)
        # Replace <br> with newline
        text = re.sub(r'<br\s*/?>', '\n', text)
        # Replace <p> tags with double newline
        text = re.sub(r'</p>\s*<p[^>]*>', '\n\n', text)
        text = re.sub(r'<p[^>]*>', '', text)
        text = re.sub(r'</p>', '', text)
        # Handle <strong> - mark for later
        text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.DOTALL)
        # Handle <code>
        text = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', text, flags=re.DOTALL)
        # Handle bullet points from styled divs
        text = re.sub(r'[•]\s*', '- ', text)
        # Remove all remaining HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Decode HTML entities
        text = html.unescape(text)
        # Restore code blocks
        for idx, code in enumerate(code_blocks):
            text = text.replace(f'[CODE_BLOCK_{idx}]', f'\n```\n{code}\n```\n')
        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        return text.strip()

    @staticmethod
    def compress_theory(atoms, max_words=400):
        """Compress atom texts to fit within word limit."""
        if not atoms:
            return atoms
        total_words = sum(len(a['text'].split()) for a in atoms)
        if total_words <= max_words:
            return atoms
        # Proportionally trim each atom
        ratio = max_words / total_words
        compressed = []
        for atom in atoms:
            words = atom['text'].split()
            target = max(15, int(len(words) * ratio))
            trimmed = ' '.join(words[:target])
            if target < len(words):
                trimmed += '...'
            compressed.append({'title': atom['title'], 'text': trimmed})
        return compressed


# ─── Exercise Selector ─────────────────────────────────────────────────────────

class ExerciseSelector:
    """Select exercises from worksheet JSON data."""

    def __init__(self, grade):
        self.grade = grade
        self.data = self._load_worksheets()

    def _load_worksheets(self):
        path = WORKSHEETS_DIR / f'cls{self.grade}.json'
        if not path.exists():
            return None
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def get_exercises_for_lesson(self, module_folder, lesson_index):
        """Get 3 exercises (1 minim, 1 standard, 1 performanta) for a lesson."""
        if not self.data:
            return []
        exercises = []
        # Try to find matching module in worksheet data
        for mod in self.data.get('modules', []):
            mod_id = mod.get('module_id', '')
            # Match by folder name similarity
            if module_folder in mod_id or mod_id in module_folder:
                lessons = mod.get('lessons', [])
                if lesson_index < len(lessons):
                    lesson = lessons[lesson_index]
                    # Pick 1 from each level
                    for level in ['minim', 'standard', 'performanta']:
                        level_data = lesson.get('levels', {}).get(level, {})
                        items = level_data.get('items', [])
                        if items:
                            exercises.append({
                                'level': level,
                                'item': items[0]  # First item from each level
                            })
                    return exercises
        return exercises

    def get_all_module_exercises(self, module_folder):
        """Get all exercises for a module folder."""
        if not self.data:
            return {}
        for mod in self.data.get('modules', []):
            mod_id = mod.get('module_id', '')
            if module_folder in mod_id or mod_id in module_folder:
                return mod
        return {}


# ─── Word Document Builder ─────────────────────────────────────────────────────

class WordDocBuilder:
    """Build Word documents with proper formatting for printable sheets."""

    def __init__(self, grade, grade_name):
        self.grade = grade
        self.grade_name = grade_name
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_page(self):
        """Set A4 page with proper margins."""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)    # 25mm for hole punch
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(f'{self.grade_name} | TIC')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

    def _setup_styles(self):
        """Create custom styles for our documents."""
        styles = self.doc.styles

        # Sheet Title style
        style = styles.add_style('SheetTitle', 1)  # 1 = paragraph
        style.font.name = 'Calibri'
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = RGBColor(30, 30, 100)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.space_before = Pt(0)

        # Module Title style
        style = styles.add_style('ModuleTitle', 1)
        style.font.name = 'Calibri'
        style.font.size = Pt(22)
        style.font.bold = True
        style.font.color.rgb = RGBColor(20, 60, 120)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section Header style
        style = styles.add_style('SectionHeader', 1)
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor(255, 255, 255)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

        # Theory text
        style = styles.add_style('TheoryText', 1)
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = Pt(14)

        # Code style
        style = styles.add_style('CodeBlock', 1)
        style.font.name = 'Consolas'
        style.font.size = Pt(9)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.space_before = Pt(2)

        # Exercise text
        style = styles.add_style('ExerciseText', 1)
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)

        # Answer line
        style = styles.add_style('AnswerLine', 1)
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(180, 180, 180)
        style.paragraph_format.space_after = Pt(2)

        # Small text for IDs and meta
        style = styles.add_style('SheetID', 1)
        style.font.name = 'Calibri'
        style.font.size = Pt(8)
        style.font.color.rgb = RGBColor(150, 150, 150)
        style.paragraph_format.space_after = Pt(0)

    def _add_section_box(self, title, color_hex='1E3A6E'):
        """Add a colored section header box."""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SectionHeader']
        # Add shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        run = p.add_run(f'  {title}')
        run.font.color.rgb = RGBColor(255, 255, 255)
        return p

    def _add_answer_lines(self, count=3):
        """Add dotted answer lines."""
        for _ in range(count):
            p = self.doc.add_paragraph()
            p.style = self.doc.styles['AnswerLine']
            p.add_run('_' * 85)

    def _add_answer_box(self, lines=6):
        """Add a bordered box for writing answers."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        # Set border
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:color="AAAAAA"/>'
            '  <w:left w:val="single" w:sz="4" w:color="AAAAAA"/>'
            '  <w:bottom w:val="single" w:sz="4" w:color="AAAAAA"/>'
            '  <w:right w:val="single" w:sz="4" w:color="AAAAAA"/>'
            '</w:tcBorders>'
        )
        tcPr.append(borders)
        # Set width
        cell.width = Cm(16)
        # Add empty lines
        for i in range(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.style = self.doc.styles['TheoryText']
            run = p.add_run(' ')
            run.font.size = Pt(10)

    def _add_mcq_options(self, options, cols=1):
        """Add MCQ checkbox options."""
        for i, opt in enumerate(options):
            letter = chr(65 + i)  # A, B, C, D
            p = self.doc.add_paragraph()
            p.style = self.doc.styles['ExerciseText']
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'\u2610 {letter}) {opt}')

    # ─── Sheet Type: Module Cover ──────────────────────────────────────────

    def add_module_cover(self, module_num, unit_name, topics, weeks, hours):
        """Add a module cover sheet."""
        sheet_id = f'CLS{self.grade}-M{module_num}'

        # Sheet ID
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetID']
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(sheet_id)

        # Add some spacing
        self.doc.add_paragraph()

        # Module title
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['ModuleTitle']
        p.add_run(f'MODUL {module_num}')

        p = self.doc.add_paragraph()
        p.style = self.doc.styles['ModuleTitle']
        run = p.add_run(unit_name)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(60, 60, 60)

        # Grade and schedule info
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{self.grade_name}  |  {weeks}  |  {hours} ore')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_paragraph()

        # Topics checklist
        self._add_section_box('CE VEI INVATA', '2D5F8A')

        for i, topic in enumerate(topics, 1):
            p = self.doc.add_paragraph()
            p.style = self.doc.styles['TheoryText']
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f'\u2610  {i}. {topic}')

        self.doc.add_paragraph()

        # Progress tracking table
        self._add_section_box('PROGRESUL MEU', '2D5F8A')

        num_lessons = min(len(topics), 9)
        table = self.doc.add_table(rows=num_lessons + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ['Lectia', 'Nota', 'Data', 'Observatii']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            # Header shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for row_idx in range(1, num_lessons + 1):
            table.cell(row_idx, 0).text = f'  L{row_idx}'
            table.cell(row_idx, 1).text = '       /10'
            table.cell(row_idx, 2).text = '     /     /2026'
            table.cell(row_idx, 3).text = ''
            for col_idx in range(4):
                for paragraph in table.cell(row_idx, col_idx).paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        self.doc.add_paragraph()

        # Bottom: module grade
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Nota modul: _____ / 10          Semnatura profesor: _________________')
        run.font.size = Pt(10)
        run.font.bold = True

        # Page break
        self.doc.add_page_break()

    # ─── Sheet Type: Theory + Exercise ─────────────────────────────────────

    def add_theory_exercise_sheet(self, module_num, lesson_num, topic_title,
                                   atoms=None, exercises=None, vocabulary=None):
        """Add a theory + exercise sheet."""
        sheet_id = f'CLS{self.grade}-M{module_num}-L{lesson_num}'

        # Sheet ID header
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetID']
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(sheet_id)

        # Title
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetTitle']
        p.add_run(f'L{lesson_num}. {topic_title}')

        # Thin separator line
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)

        # ── THEORY SECTION ──
        self._add_section_box('TEORIE', '1E3A6E')

        if atoms:
            for atom in atoms:
                # Atom title
                p = self.doc.add_paragraph()
                p.style = self.doc.styles['TheoryText']
                run = p.add_run(f'{atom["title"]}')
                run.bold = True
                run.font.size = Pt(10)

                # Atom text - split by code blocks first
                text = atom['text']
                # Split on triple-backtick code blocks
                code_split = re.split(r'```\n?(.*?)\n?```', text, flags=re.DOTALL)
                for seg_idx, segment in enumerate(code_split):
                    if seg_idx % 2 == 1:
                        # This is a code block - render in monospace with background
                        for code_line in segment.strip().split('\n'):
                            p = self.doc.add_paragraph()
                            p.style = self.doc.styles['CodeBlock']
                            p.paragraph_format.left_indent = Cm(0.5)
                            # Add light gray background
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
                            p.paragraph_format.element.get_or_add_pPr().append(shading)
                            run = p.add_run(f'  {code_line}')
                    else:
                        # Normal text - handle bold markers and inline code
                        segment = segment.strip()
                        if not segment:
                            continue
                        # Handle inline `code` separately
                        inline_parts = re.split(r'`([^`]+)`', segment)
                        p = self.doc.add_paragraph()
                        p.style = self.doc.styles['TheoryText']
                        for ip_idx, ip in enumerate(inline_parts):
                            if ip_idx % 2 == 1:
                                # Inline code
                                run = p.add_run(ip)
                                run.font.name = 'Consolas'
                                run.font.size = Pt(9)
                            else:
                                # Normal text with bold markers
                                bold_parts = re.split(r'\*\*(.*?)\*\*', ip)
                                for bp_idx, bp in enumerate(bold_parts):
                                    if not bp:
                                        continue
                                    run = p.add_run(bp)
                                    if bp_idx % 2 == 1:
                                        run.bold = True
        else:
            # No theory available - add placeholder
            p = self.doc.add_paragraph()
            p.style = self.doc.styles['TheoryText']
            run = p.add_run(f'[Tema: {topic_title}]')
            run.italic = True
            self._add_answer_lines(4)

        # ── VOCABULARY SECTION (if provided) ──
        if vocabulary and len(vocabulary) > 0:
            self._add_section_box('VOCABULAR', '4A7C59')
            table = self.doc.add_table(rows=len(vocabulary) + 1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Header
            table.cell(0, 0).text = 'Termen'
            table.cell(0, 1).text = 'Definitie'
            for i in range(2):
                for run in table.cell(0, i).paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0E8" w:val="clear"/>')
                table.cell(0, i)._tc.get_or_add_tcPr().append(shading)
            # Terms
            for idx, (term, defn) in enumerate(vocabulary, 1):
                table.cell(idx, 0).text = term
                table.cell(idx, 1).text = defn
                for col in range(2):
                    for paragraph in table.cell(idx, col).paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

        # ── EXERCISES SECTION ──
        self._add_section_box('EXERCITII', '8B4513')

        if exercises:
            for i, ex in enumerate(exercises, 1):
                level = ex.get('level', '')
                item = ex.get('item', {})
                level_label = {'minim': 'Baza', 'standard': 'Standard', 'performanta': 'Avansat'}.get(level, level)
                level_stars = {'minim': '*', 'standard': '**', 'performanta': '***'}.get(level, '')

                ex_type = item.get('type', 'short')
                question = item.get('question', item.get('description', ''))

                # Exercise header
                p = self.doc.add_paragraph()
                p.style = self.doc.styles['ExerciseText']
                run = p.add_run(f'E{i}. ({level_label} {level_stars}) ')
                run.bold = True
                run.font.size = Pt(10)

                if ex_type == 'mcq':
                    run = p.add_run(question)
                    run.font.size = Pt(10)
                    options = item.get('options', [])
                    self._add_mcq_options(options)
                elif ex_type in ('short', 'explain'):
                    run = p.add_run(question)
                    run.font.size = Pt(10)
                    lines = 2 if ex_type == 'short' else 4
                    self._add_answer_lines(lines)
                elif ex_type in ('task', 'create'):
                    run = p.add_run(question)
                    run.font.size = Pt(10)
                    # Show criteria if available
                    criteria = item.get('criteria', item.get('requirements', []))
                    if criteria:
                        p = self.doc.add_paragraph()
                        p.style = self.doc.styles['ExerciseText']
                        p.paragraph_format.left_indent = Cm(0.5)
                        run = p.add_run('Criterii: ' + ', '.join(criteria))
                        run.font.size = Pt(9)
                        run.italic = True
                    self._add_answer_box(5)
                elif ex_type == 'debug':
                    run = p.add_run(item.get('question', question))
                    run.font.size = Pt(10)
                    if item.get('broken_code'):
                        p = self.doc.add_paragraph()
                        p.style = self.doc.styles['CodeBlock']
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.add_run(item['broken_code'])
                    self._add_answer_lines(3)
                elif ex_type == 'scenario':
                    context = item.get('context', '')
                    problem = item.get('problem', '')
                    q = item.get('question', question)
                    run = p.add_run(f'{context} {problem}')
                    run.font.size = Pt(10)
                    p2 = self.doc.add_paragraph()
                    p2.style = self.doc.styles['ExerciseText']
                    p2.paragraph_format.left_indent = Cm(0.5)
                    run2 = p2.add_run(q)
                    run2.font.size = Pt(10)
                    run2.italic = True
                    self._add_answer_lines(4)
                else:
                    run = p.add_run(question)
                    run.font.size = Pt(10)
                    self._add_answer_lines(3)

                # Small space between exercises
                self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            # No exercises from worksheet - add generic framework
            for i, (level, label, stars) in enumerate(
                [('minim', 'Baza', '*'), ('standard', 'Standard', '**'), ('performanta', 'Avansat', '***')], 1
            ):
                p = self.doc.add_paragraph()
                p.style = self.doc.styles['ExerciseText']
                run = p.add_run(f'E{i}. ({label} {stars}) ')
                run.bold = True
                run = p.add_run(f'[Exercitiu despre: {topic_title}]')
                run.italic = True
                self._add_answer_lines(3)

        # ── NOTES SECTION ──
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['TheoryText']
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run('Notite:')
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        self._add_answer_lines(2)

        # Page break
        self.doc.add_page_break()

    # ─── Sheet Type: Reference Card ────────────────────────────────────────

    def add_reference_card(self, title, sections):
        """Add a reference card sheet.
        sections: list of {'title': str, 'items': list of (term, desc) or list of str}
        """
        sheet_id = f'CLS{self.grade}-REF'

        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetID']
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(sheet_id)

        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetTitle']
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'FISA DE REFERINTA: {title}')
        run.font.size = Pt(14)

        for section in sections:
            self._add_section_box(section['title'], '2D5F8A')
            items = section.get('items', [])

            if items and isinstance(items[0], (tuple, list)):
                # Table format (term, description)
                table = self.doc.add_table(rows=len(items), cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for idx, (term, desc) in enumerate(items):
                    table.cell(idx, 0).text = term
                    table.cell(idx, 1).text = desc
                    for col in range(2):
                        for paragraph in table.cell(idx, col).paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                            if col == 0:
                                for run in paragraph.runs:
                                    run.font.bold = True
            else:
                # List format
                for item in items:
                    p = self.doc.add_paragraph()
                    p.style = self.doc.styles['TheoryText']
                    p.paragraph_format.left_indent = Cm(0.3)
                    run = p.add_run(f'- {item}')
                    run.font.size = Pt(9)

        self.doc.add_page_break()

    # ─── Sheet Type: Assessment ────────────────────────────────────────────

    def add_assessment_sheet(self, module_num, unit_name, questions):
        """Add an assessment sheet.
        questions: list of {'part': str, 'text': str, 'points': float, 'type': 'mcq'|'short'|'code'|'open', 'options': []}
        """
        sheet_id = f'CLS{self.grade}-M{module_num}-TEST'

        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetID']
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(sheet_id)

        p = self.doc.add_paragraph()
        p.style = self.doc.styles['SheetTitle']
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'TEST RECAPITULATIV - Modul {module_num}')
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{self.grade_name}  |  {unit_name}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Nume: ____________________  Clasa: _______  Data: ___/___/2026')
        run.font.size = Pt(10)

        self.doc.add_paragraph()

        # Parts: I = Knowledge (4p), II = Application (3p), III = Problem solving (3p)
        current_part = None
        q_num = 0

        for q in questions:
            part = q.get('part', '')
            if part != current_part:
                current_part = part
                part_titles = {
                    'I': 'PARTEA I - Cunostinte (4p)',
                    'II': 'PARTEA II - Aplicare (3p)',
                    'III': 'PARTEA III - Rezolvare (3p)'
                }
                self._add_section_box(part_titles.get(part, part), '5A2D82')

            q_num += 1
            q_type = q.get('type', 'short')
            text = q.get('text', '')
            points = q.get('points', 1)

            p = self.doc.add_paragraph()
            p.style = self.doc.styles['ExerciseText']
            run = p.add_run(f'{q_num}. ')
            run.bold = True
            run = p.add_run(f'{text} ')
            run = p.add_run(f'({points}p)')
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(120, 120, 120)

            if q_type == 'mcq':
                options = q.get('options', [])
                self._add_mcq_options(options)
            elif q_type == 'short':
                self._add_answer_lines(2)
            elif q_type == 'code':
                self._add_answer_box(6)
            elif q_type == 'open':
                self._add_answer_lines(4)

        # Bottom scoring
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('1p din oficiu + _____ p = Nota: _____ / 10')
        run.font.size = Pt(11)
        run.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Semnatura profesor: _______________     Data: ___/___/2026')
        run.font.size = Pt(10)

        self.doc.add_page_break()

    def save(self, filepath):
        """Save the document."""
        self.doc.save(str(filepath))


# ─── Reference Card Data ──────────────────────────────────────────────────────

REFERENCE_CARDS = {
    5: [
        {
            'title': 'Scurtaturi de Tastatura',
            'sections': [
                {'title': 'GENERALE', 'items': [
                    ('Ctrl+C', 'Copiere'), ('Ctrl+V', 'Lipire'), ('Ctrl+X', 'Decupare'),
                    ('Ctrl+Z', 'Anulare (Undo)'), ('Ctrl+Y', 'Refacere (Redo)'),
                    ('Ctrl+S', 'Salvare'), ('Ctrl+A', 'Selecteaza tot'),
                    ('Ctrl+P', 'Tiparire'), ('Alt+F4', 'Inchide aplicatia'),
                    ('F2', 'Redenumire fisier'), ('Delete', 'Sterge'),
                    ('Windows+E', 'Deschide File Explorer'),
                ]},
                {'title': 'NAVIGARE', 'items': [
                    ('Tab', 'Urmatorul camp/element'), ('Alt+Tab', 'Comuta intre aplicatii'),
                    ('Windows+D', 'Afiseaza desktop-ul'),
                ]}
            ]
        },
        {
            'title': 'Operatii cu Fisiere si Foldere',
            'sections': [
                {'title': 'OPERATII FISIERE', 'items': [
                    ('Creare', 'Click dreapta > New > tip fisier'),
                    ('Copiere', 'Ctrl+C apoi Ctrl+V in alta locatie'),
                    ('Mutare', 'Ctrl+X apoi Ctrl+V sau drag & drop'),
                    ('Redenumire', 'F2 sau click dreapta > Rename'),
                    ('Stergere', 'Delete (Recycle Bin) sau Shift+Delete (permanent)'),
                ]},
                {'title': 'EXTENSII COMUNE', 'items': [
                    ('.docx', 'Document Word'), ('.xlsx', 'Foaie Excel'),
                    ('.pptx', 'Prezentare PowerPoint'), ('.pdf', 'Document portabil'),
                    ('.jpg/.png', 'Imagine'), ('.mp3/.mp4', 'Audio/Video'),
                    ('.txt', 'Text simplu'), ('.html', 'Pagina web'),
                ]}
            ]
        },
        {
            'title': 'Blocuri Scratch',
            'sections': [
                {'title': 'MISCARE (albastru)', 'items': [
                    ('mergi N pasi', 'Deplaseaza sprite-ul inainte'),
                    ('roteste N grade', 'Roteste la dreapta/stanga'),
                    ('mergi la x: y:', 'Pozitioneaza exact'),
                ]},
                {'title': 'ASPECT (violet)', 'items': [
                    ('spune "text"', 'Afiseaza balon text'),
                    ('schimba costumul', 'Animatie sprite'),
                    ('arata/ascunde', 'Vizibilitate sprite'),
                ]},
                {'title': 'CONTROL (galben)', 'items': [
                    ('cand flag verde', 'Start program'),
                    ('repeta N ori', 'Bucla numarata'),
                    ('daca...atunci', 'Conditie (decizie)'),
                    ('repeta pana cand', 'Bucla conditionata'),
                ]},
                {'title': 'VARIABILE (portocaliu)', 'items': [
                    ('seteaza var la N', 'Atribuie valoare'),
                    ('modifica var cu N', 'Aduna/scade valoare'),
                ]}
            ]
        }
    ],
    6: [
        {
            'title': 'PowerPoint - Scurtaturi si Sfaturi',
            'sections': [
                {'title': 'CREARE PREZENTARE', 'items': [
                    ('Ctrl+M', 'Slide nou'), ('Ctrl+D', 'Duplica slide-ul'),
                    ('F5', 'Porneste prezentarea'), ('Esc', 'Opreste prezentarea'),
                    ('Ctrl+G', 'Grupeaza obiecte'),
                ]},
                {'title': 'REGULI PREZENTARE', 'items': [
                    'Maxim 6 cuvinte pe rand, 6 randuri pe slide (Regula 6x6)',
                    'Font minim 24pt pentru public',
                    'Contrast mare: text inchis pe fundal deschis (sau invers)',
                    'O idee principala per slide',
                    'Imagini relevante, nu decorative',
                ]}
            ]
        },
        {
            'title': 'Reguli Email si Comunicare Online',
            'sections': [
                {'title': 'STRUCTURA EMAIL', 'items': [
                    ('To/Catre', 'Destinatarul principal'),
                    ('CC', 'Copie carbon - informare'),
                    ('BCC', 'Copie ascunsa'),
                    ('Subject/Subiect', 'Rezumat clar al mesajului'),
                    ('Atasament', 'Fisiere atasate (atentie la dimensiune!)'),
                ]},
                {'title': 'NETICHETA', 'items': [
                    'Scrie un subiect clar si descriptiv',
                    'Saluta politicos (Buna ziua, Stimate/a...)',
                    'Nu scrie TOT CU MAJUSCULE (= tipat)',
                    'Reciteste inainte de a trimite',
                    'Nu trimite informatii personale pe email',
                    'Raspunde in maxim 24 ore',
                ]}
            ]
        },
        {
            'title': 'Scratch - Blocuri de Control',
            'sections': [
                {'title': 'STRUCTURI DE CONTROL', 'items': [
                    ('repeta N ori', 'FOR: executa de N ori'),
                    ('repeta la infinit', 'WHILE TRUE: nu se opreste'),
                    ('repeta pana cand <cond>', 'DO-WHILE: pana cand conditia e adevarata'),
                    ('daca <cond> atunci', 'IF: executa doar daca e adevarat'),
                    ('daca <cond> atunci / altfel', 'IF-ELSE: doua ramuri'),
                ]},
                {'title': 'OPERATORI', 'items': [
                    ('= < >', 'Comparatie (egal, mai mic, mai mare)'),
                    ('si / sau / nu', 'Operatori logici'),
                    ('+ - * /', 'Operatori aritmetici'),
                    ('rest la impartirea / rotunjire', 'Operatii matematice speciale'),
                ]}
            ]
        }
    ],
    7: [
        {
            'title': 'Microsoft Word - Toolbar Formatare',
            'sections': [
                {'title': 'FORMATARE TEXT', 'items': [
                    ('Ctrl+B', 'Bold (ingrosat)'), ('Ctrl+I', 'Italic (inclinat)'),
                    ('Ctrl+U', 'Underline (subliniat)'), ('Ctrl+E', 'Aliniere centru'),
                    ('Ctrl+L', 'Aliniere stanga'), ('Ctrl+R', 'Aliniere dreapta'),
                    ('Ctrl+J', 'Aliniere justified (stanga-dreapta)'),
                ]},
                {'title': 'FORMATARE PAGINA', 'items': [
                    ('Antet (Header)', 'Insert > Header - text repetat sus'),
                    ('Subsol (Footer)', 'Insert > Footer - text repetat jos'),
                    ('Nr. pagina', 'Insert > Page Number'),
                    ('Orientare', 'Layout > Orientation (Portrait/Landscape)'),
                    ('Margini', 'Layout > Margins (Normal = 2.54 cm)'),
                ]},
                {'title': 'REGULI TEHNOREDACTARE', 'items': [
                    'Dupa virgula, punct, punct si virgula: un spatiu',
                    'Inainte de virgula, punct: FARA spatiu',
                    'Un singur spatiu intre cuvinte (nu mai multe)',
                    'Paragraf nou = Enter; Rand nou = Shift+Enter',
                    'Diacritice obligatorii: a, i, s, t (RO keyboard)',
                ]}
            ]
        },
        {
            'title': 'Formate Audio-Video',
            'sections': [
                {'title': 'FORMATE AUDIO', 'items': [
                    ('.mp3', 'Cel mai comun, comprimat cu pierderi'),
                    ('.wav', 'Necomprimat, calitate maxima, fisier mare'),
                    ('.aac', 'Calitate buna, folosit de Apple/YouTube'),
                    ('.ogg', 'Open source, calitate buna'),
                ]},
                {'title': 'FORMATE VIDEO', 'items': [
                    ('.mp4', 'Universal, compatibil cu toate dispozitivele'),
                    ('.avi', 'Vechi, fisiere mari, calitate variabila'),
                    ('.mkv', 'Suporta multiple piste audio/subtitrari'),
                    ('.mov', 'Format Apple, calitate buna'),
                ]},
                {'title': 'CONCEPTE EDITARE', 'items': [
                    ('Timeline', 'Linia de timp - unde aranjezi clipurile'),
                    ('Tranzitie', 'Efect intre doua clipuri (dissolve, fade)'),
                    ('Rendering', 'Export final - procesarea proiectului'),
                    ('FPS', 'Cadre pe secunda (24=cinema, 30=TV, 60=gaming)'),
                ]}
            ]
        },
        {
            'title': 'Python - Sintaxa de Baza',
            'sections': [
                {'title': 'TIPURI DE DATE', 'items': [
                    ('int', 'Numar intreg: 42, -7, 0'),
                    ('float', 'Numar real: 3.14, -2.5'),
                    ('str', 'Text (string): "Salut", \'Ana\''),
                    ('bool', 'Adevarat/Fals: True, False'),
                ]},
                {'title': 'STRUCTURI', 'items': [
                    ('x = 5', 'Atribuire variabila'),
                    ('x = input("Text: ")', 'Citire de la tastatura'),
                    ('print(x)', 'Afisare pe ecran'),
                    ('if x > 0:', 'Conditie (daca)'),
                    ('elif x == 0:', 'Altfel daca'),
                    ('else:', 'Altfel'),
                    ('for i in range(n):', 'Bucla numarata'),
                    ('while cond:', 'Bucla conditionata'),
                ]},
                {'title': 'OPERATORI', 'items': [
                    ('+ - * /', 'Aritmetici (/ = impartire reala)'),
                    ('//', 'Impartire intreaga (fara rest)'),
                    ('%', 'Modulo (restul impartirii)'),
                    ('**', 'Putere (2**3 = 8)'),
                    ('== != < > <= >=', 'Comparatie'),
                    ('and or not', 'Logici'),
                ]}
            ]
        },
        {
            'title': 'Aplicatii Colaborative',
            'sections': [
                {'title': 'GOOGLE WORKSPACE', 'items': [
                    ('Google Docs', 'Editor text online, colaborare in timp real'),
                    ('Google Sheets', 'Calcul tabelar online'),
                    ('Google Slides', 'Prezentari online'),
                    ('Google Drive', 'Stocare in cloud, partajare fisiere'),
                ]},
                {'title': 'ALTE INSTRUMENTE', 'items': [
                    ('Padlet', 'Tabla virtuala colaborativa'),
                    ('Canva', 'Design grafic simplu, template-uri'),
                    ('Microsoft 365', 'Varianta online Word/Excel/PowerPoint'),
                ]},
                {'title': 'ETICA DIGITALA', 'items': [
                    'Nu copia munca altora fara sa dai credit (plagiat)',
                    'Protejeaza datele personale (nu le publica online)',
                    'Respecta drepturile de autor (imagini, texte, muzica)',
                    'Verifica informatiile inainte de a le distribui',
                ]}
            ]
        }
    ],
    8: [
        {
            'title': 'Excel - Functii si Formule',
            'sections': [
                {'title': 'FUNCTII DE BAZA', 'items': [
                    ('=SUM(A1:A10)', 'Suma valorilor din interval'),
                    ('=AVERAGE(A1:A10)', 'Media aritmetica'),
                    ('=MIN(A1:A10)', 'Valoarea minima'),
                    ('=MAX(A1:A10)', 'Valoarea maxima'),
                    ('=COUNT(A1:A10)', 'Numara celulele cu numere'),
                    ('=COUNTA(A1:A10)', 'Numara celulele nevide'),
                    ('=IF(cond, da, nu)', 'Functia conditionala'),
                ]},
                {'title': 'REFERINTE CELULE', 'items': [
                    ('A1', 'Referinta relativa (se modifica la copiere)'),
                    ('$A$1', 'Referinta absoluta (ramane fixa)'),
                    ('$A1 / A$1', 'Referinta mixta (o directie fixa)'),
                ]},
                {'title': 'TIPURI DIAGRAME', 'items': [
                    ('Coloana/Bara', 'Comparatii intre categorii'),
                    ('Linie', 'Evolutie in timp'),
                    ('Cerc (Pie)', 'Proportii din total (100%)'),
                    ('Scatter (XY)', 'Relatia intre doua variabile'),
                ]}
            ]
        },
        {
            'title': 'HTML si CSS - Etichete de Baza',
            'sections': [
                {'title': 'STRUCTURA HTML', 'items': [
                    ('<html>', 'Elementul radacina'),
                    ('<head>', 'Metadate (title, CSS, meta)'),
                    ('<body>', 'Continutul vizibil al paginii'),
                    ('<h1>...<h6>', 'Titluri (h1=mare, h6=mic)'),
                    ('<p>', 'Paragraf de text'),
                    ('<a href="url">', 'Link (hyperlink)'),
                    ('<img src="poza.jpg">', 'Imagine'),
                    ('<ul> / <ol> / <li>', 'Lista neordonata / ordonata / element'),
                    ('<table> <tr> <td>', 'Tabel / rand / celula'),
                    ('<br>', 'Rand nou (fara paragraf)'),
                    ('<strong> / <em>', 'Bold / Italic'),
                ]},
                {'title': 'CSS DE BAZA', 'items': [
                    ('color: red;', 'Culoarea textului'),
                    ('background-color: blue;', 'Culoarea fundalului'),
                    ('font-size: 16px;', 'Dimensiunea fontului'),
                    ('text-align: center;', 'Alinierea textului'),
                    ('margin: 10px;', 'Spatiu exterior'),
                    ('padding: 10px;', 'Spatiu interior'),
                    ('border: 1px solid;', 'Bordura'),
                ]}
            ]
        },
        {
            'title': 'C++ - Sintaxa de Baza',
            'sections': [
                {'title': 'STRUCTURA PROGRAM', 'items': [
                    ('#include <iostream>', 'Biblioteca intrare/iesire'),
                    ('using namespace std;', 'Spatiu de nume standard'),
                    ('int main() { }', 'Functia principala'),
                    ('cout << "text";', 'Afisare pe ecran'),
                    ('cin >> x;', 'Citire de la tastatura'),
                    ('return 0;', 'Program terminat cu succes'),
                ]},
                {'title': 'TIPURI SI VARIABILE', 'items': [
                    ('int x = 5;', 'Numar intreg'),
                    ('float y = 3.14;', 'Numar real'),
                    ('char c = \'A\';', 'Un singur caracter'),
                    ('string s = "text";', 'Sir de caractere'),
                    ('int a[5];', 'Tablou (array) de 5 elemente'),
                ]},
                {'title': 'STRUCTURI DE CONTROL', 'items': [
                    ('if (cond) { }', 'Decizie simpla'),
                    ('if...else if...else', 'Decizie multipla'),
                    ('for (i=0; i<n; i++)', 'Bucla numarata'),
                    ('while (cond) { }', 'Bucla conditionata anterior'),
                    ('do { } while (cond);', 'Bucla conditionata posterior'),
                ]},
                {'title': 'TABLOURI (ARRAYS)', 'items': [
                    ('int a[5] = {1,2,3,4,5};', 'Declarare cu initializare'),
                    ('a[0]', 'Primul element (index de la 0!)'),
                    ('a[n-1]', 'Ultimul element'),
                    ('for(i=0;i<n;i++) cout<<a[i];', 'Parcurgere tablou'),
                ]}
            ]
        },
        {
            'title': 'Robot Didactic - Algoritmi',
            'sections': [
                {'title': 'CONCEPTE ROBOTICA', 'items': [
                    ('Senzor', 'Dispozitiv care masoara mediul (distanta, culoare, lumina)'),
                    ('Actuator', 'Dispozitiv care executa actiuni (motor, LED, sunet)'),
                    ('Algoritm', 'Secventa de pasi pentru rezolvarea unei probleme'),
                ]},
                {'title': 'ALGORITMI DE BAZA', 'items': [
                    ('Detectare obstacol', 'Citeste senzor distanta > Daca < prag > Opreste + Roteste'),
                    ('Urmarire linie', 'Citeste senzor culoare > Daca negru: inainte > Daca alb: corecteaza'),
                    ('Parcurgere traseu', 'Combina detectare + urmarire + numarare intersectii'),
                ]}
            ]
        }
    ]
}

# ─── Assessment Templates ─────────────────────────────────────────────────────

def generate_assessment_questions(grade, module_num, topics):
    """Generate assessment questions from topic list."""
    questions = []

    # Part I: Knowledge (4 points) - 2-3 questions
    if len(topics) >= 1:
        questions.append({
            'part': 'I', 'type': 'short', 'points': 2,
            'text': f'Defineste (in 1-2 propozitii) notiunea de: {topics[0].split(".")[0].split("-")[0].strip()}'
        })
    if len(topics) >= 3:
        # MCQ from middle topics
        questions.append({
            'part': 'I', 'type': 'mcq', 'points': 1,
            'text': f'Care dintre urmatoarele NU este legat de tema "{topics[1]}"?',
            'options': ['[varianta A]', '[varianta B]', '[varianta C - corecta]', '[varianta D]']
        })
    questions.append({
        'part': 'I', 'type': 'short', 'points': 1,
        'text': f'Enumera 3 concepte cheie din modulul "{topics[0]}" (fara explicatii).'
    })

    # Part II: Application (3 points) - 1-2 questions
    questions.append({
        'part': 'II', 'type': 'short', 'points': 1.5,
        'text': f'Explica cum se aplica in practica: {topics[min(2, len(topics)-1)]}'
    })
    questions.append({
        'part': 'II', 'type': 'short', 'points': 1.5,
        'text': f'Gaseste si explica eroarea (daca exista) in exemplul de mai jos, legat de {topics[min(1, len(topics)-1)]}.'
    })

    # Part III: Problem solving (3 points) - 1 question
    questions.append({
        'part': 'III', 'type': 'code', 'points': 3,
        'text': f'Rezolva: Folosind cunostintele din acest modul, realizeaza exercitiul practic referitor la {topics[-1]}.'
    })

    return questions


# ─── Main Generator ───────────────────────────────────────────────────────────

def load_content_map():
    with open(CONTENT_MAP, 'r', encoding='utf-8') as f:
        return json.load(f)

def find_lesson_files(grade, folder_name):
    """Find all lesson HTML files in a content folder."""
    folder_path = CONTENT_DIR / f'cls{grade}' / folder_name
    if not folder_path.exists():
        return []
    files = sorted([f for f in folder_path.iterdir()
                    if f.name.startswith('lectia') and f.name.endswith('.html')])
    return files

def extract_theory_from_file(filepath):
    """Extract theory atoms and quiz questions from a single lesson file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        atoms = TheoryExtractor.extract_atoms(content)
        goal = TheoryExtractor.extract_goal(content)
        title = TheoryExtractor.extract_title(content)
        quizzes = TheoryExtractor.extract_quizzes(content)
        return {
            'atoms': TheoryExtractor.compress_theory(atoms),
            'goal': goal,
            'title': title,
            'quizzes': quizzes
        }
    except Exception as e:
        print(f'  WARNING: Could not extract theory from {filepath}: {e}')
        return {'atoms': [], 'goal': '', 'title': '', 'quizzes': []}


def generate_grade(grade, content_map, target_module=None):
    """Generate all printable sheets for a grade."""
    grade_str = str(grade)
    grade_data = content_map['grades'].get(grade_str, {})
    if not grade_data:
        print(f'ERROR: No data for grade {grade}')
        return

    grade_names = {5: 'Clasa a V-a', 6: 'Clasa a VI-a', 7: 'Clasa a VII-a', 8: 'Clasa a VIII-a'}
    grade_name = grade_names.get(grade, f'Clasa {grade}')

    exercise_selector = ExerciseSelector(grade)

    modules = grade_data.get('modules', {})

    for mod_key in sorted(modules.keys()):
        mod_num = int(mod_key[1:])  # M1 -> 1
        if target_module and mod_num != target_module:
            continue

        mod_data = modules[mod_key]
        unit_name = mod_data.get('unit_name', f'Modul {mod_num}')
        topics = mod_data.get('topics', [])
        content_paths = mod_data.get('content_paths', [])
        hours = mod_data.get('hours', 0)
        weeks = mod_data.get('weeks', '')

        print(f'\n  Generating CLS{grade}-M{mod_num}: {unit_name} ({len(topics)} topics)')

        doc = WordDocBuilder(grade, grade_name)

        # 1. MODULE COVER
        doc.add_module_cover(mod_num, unit_name, topics, weeks, hours)

        # 2. THEORY + EXERCISE SHEETS
        # Collect all lesson files from all content_paths
        all_lesson_files = []
        for folder in content_paths:
            files = find_lesson_files(grade, folder)
            all_lesson_files.extend(files)

        for t_idx, topic in enumerate(topics):
            lesson_num = t_idx + 1

            # Try to match with a lesson file
            atoms = []
            quizzes = []
            if t_idx < len(all_lesson_files):
                theory = extract_theory_from_file(all_lesson_files[t_idx])
                atoms = theory.get('atoms', [])
                quizzes = theory.get('quizzes', [])

            # Try to get exercises from worksheet JSON first
            exercises = []
            for folder in content_paths:
                exercises = exercise_selector.get_exercises_for_lesson(folder, t_idx)
                if exercises:
                    break

            # If no worksheet exercises, build from HTML quiz data
            if not exercises and quizzes:
                # Use up to 3 quiz questions as MCQ exercises
                for q_idx, quiz in enumerate(quizzes[:3]):
                    level = ['minim', 'standard', 'performanta'][min(q_idx, 2)]
                    exercises.append({
                        'level': level,
                        'item': {
                            'type': 'mcq',
                            'question': quiz['question'],
                            'options': quiz['options']
                        }
                    })
                # If we have quiz MCQs, add a writing exercise too
                if len(exercises) < 3:
                    exercises.append({
                        'level': 'performanta',
                        'item': {
                            'type': 'explain',
                            'question': f'Explica cu cuvintele tale conceptul principal din aceasta lectie ({topic}).'
                        }
                    })

            # Extract vocabulary from atoms (bold terms with context as definition)
            vocabulary = []
            for atom in atoms:
                text = atom.get('text', '')
                bold_terms = re.findall(r'\*\*(.*?)\*\*', text)
                for term in bold_terms[:3]:  # Max 3 per atom
                    if len(term) < 50 and len(term) > 2:  # Skip long phrases and tiny ones
                        # Extract definition: sentence containing the bold term
                        clean_text = text.replace(f'**{term}**', term)
                        sentences = re.split(r'[.!?]\s+', clean_text)
                        defn = ''
                        for sent in sentences:
                            if term.lower() in sent.lower() and len(sent) > len(term) + 5:
                                defn = sent.strip().rstrip('.')
                                if len(defn) > 80:
                                    defn = defn[:77] + '...'
                                break
                        vocabulary.append((term, defn))

            # Deduplicate vocabulary
            seen = set()
            unique_vocab = []
            for term, defn in vocabulary:
                if term.lower() not in seen:
                    seen.add(term.lower())
                    unique_vocab.append((term, defn))
            vocabulary = unique_vocab[:6]  # Max 6 terms total

            doc.add_theory_exercise_sheet(
                module_num=mod_num,
                lesson_num=lesson_num,
                topic_title=topic,
                atoms=atoms if atoms else None,
                exercises=exercises if exercises else None,
                vocabulary=vocabulary if vocabulary else None
            )

        # 3. ASSESSMENT SHEET
        assessment_questions = generate_assessment_questions(grade, mod_num, topics)
        doc.add_assessment_sheet(mod_num, unit_name, assessment_questions)

        # Save module document
        slug = unit_name.lower().replace(' ', '-').replace('/', '-')[:30]
        slug = re.sub(r'[^a-z0-9-]', '', slug)
        filename = f'cls{grade}_M{mod_num}_{slug}.docx'
        output_path = OUTPUT_DIR / filename
        doc.save(output_path)
        print(f'    -> Saved: {filename}')

    # 4. REFERENCE CARDS (one file per grade, unless targeting single module)
    if not target_module:
        ref_cards = REFERENCE_CARDS.get(grade, [])
        if ref_cards:
            print(f'\n  Generating CLS{grade} reference cards ({len(ref_cards)} cards)')
            ref_doc = WordDocBuilder(grade, grade_name)
            for card in ref_cards:
                ref_doc.add_reference_card(card['title'], card['sections'])
            ref_filename = f'cls{grade}_reference_cards.docx'
            ref_doc.save(OUTPUT_DIR / ref_filename)
            print(f'    -> Saved: {ref_filename}')


def main():
    parser = argparse.ArgumentParser(description='Generate printable ICT worksheets')
    parser.add_argument('--grade', type=int, choices=[5, 6, 7, 8], help='Generate for specific grade')
    parser.add_argument('--module', type=int, choices=[1, 2, 3, 4, 5], help='Generate for specific module')
    args = parser.parse_args()

    print('=' * 60)
    print('  PRINTABLE SHEETS GENERATOR - Foi Volante TIC')
    print('=' * 60)

    content_map = load_content_map()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    grades = [args.grade] if args.grade else [5, 6, 7, 8]

    for grade in grades:
        print(f'\n{"="*40}')
        print(f'  GRADE: Clasa a {grade}-a')
        print(f'{"="*40}')
        generate_grade(grade, content_map, target_module=args.module)

    print(f'\n{"="*60}')
    print(f'  DONE! Output: {OUTPUT_DIR}')
    print(f'{"="*60}')


if __name__ == '__main__':
    main()
